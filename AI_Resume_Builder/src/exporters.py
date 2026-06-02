"""
exporters.py — Resume export to PDF, DOCX, and TXT formats.

Uses ReportLab for PDF generation and python-docx for Word documents.
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.config import DOCX_FONT_NAME, DOCX_FONT_SIZE_BODY, PDF_MARGIN_INCH
from src.models import Resume

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────
# PDF Export
# ──────────────────────────────────────────────────────────────


def export_to_pdf(resume: Resume, output_path: Path) -> None:
    """
    Export resume to PDF using ReportLab.

    Args:
        resume: Resume data.
        output_path: Destination file path.
    """
    logger.info(f"Exporting resume to PDF: {output_path}")

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=PDF_MARGIN_INCH * inch,
        rightMargin=PDF_MARGIN_INCH * inch,
        topMargin=PDF_MARGIN_INCH * inch,
        bottomMargin=PDF_MARGIN_INCH * inch,
    )

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=colors.HexColor("#2C3E50"),
        spaceAfter=6,
        alignment=1,  # Center
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        textColor=colors.HexColor("#34495E"),
        spaceAfter=6,
        spaceBefore=12,
        borderColor=colors.HexColor("#3498DB"),
        borderWidth=0,
        borderPadding=3,
    )
    body_style = styles["BodyText"]

    # Personal Info
    pi = resume.personal_info
    story.append(Paragraph(pi.full_name, title_style))
    contact_info = f"{pi.email} | {pi.phone} | {pi.location}"
    if pi.linkedin:
        contact_info += f" | LinkedIn: {pi.linkedin}"
    story.append(Paragraph(contact_info, body_style))
    story.append(Spacer(1, 0.2 * inch))

    # Summary
    if pi.summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        story.append(Paragraph(pi.summary, body_style))
        story.append(Spacer(1, 0.15 * inch))

    # Experience
    if resume.experience:
        story.append(Paragraph("EXPERIENCE", heading_style))
        for exp in resume.experience:
            story.append(Paragraph(
                f"<b>{exp.position}</b> at {exp.company}",
                body_style
            ))
            story.append(Paragraph(
                f"{exp.location} | {exp.start_date} - {exp.end_date}",
                body_style
            ))
            for item in exp.responsibilities + exp.achievements:
                story.append(Paragraph(f"• {item}", body_style))
            story.append(Spacer(1, 0.1 * inch))

    # Education
    if resume.education:
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in resume.education:
            story.append(Paragraph(
                f"<b>{edu.degree} in {edu.field}</b>",
                body_style
            ))
            story.append(Paragraph(
                f"{edu.institution} | {edu.start_date} - {edu.end_date}",
                body_style
            ))
            if edu.gpa:
                story.append(Paragraph(f"GPA: {edu.gpa}", body_style))
            story.append(Spacer(1, 0.1 * inch))

    # Skills
    if resume.skills:
        story.append(Paragraph("SKILLS", heading_style))
        skills_text = ", ".join(resume.skills)
        story.append(Paragraph(skills_text, body_style))
        story.append(Spacer(1, 0.1 * inch))

    # Projects
    if resume.projects:
        story.append(Paragraph("PROJECTS", heading_style))
        for proj in resume.projects:
            story.append(Paragraph(f"<b>{proj.title}</b>", body_style))
            story.append(Paragraph(proj.description, body_style))
            if proj.technologies:
                story.append(Paragraph(
                    f"Technologies: {', '.join(proj.technologies)}",
                    body_style
                ))
            story.append(Spacer(1, 0.1 * inch))

    # Certifications
    if resume.certifications:
        story.append(Paragraph("CERTIFICATIONS", heading_style))
        for cert in resume.certifications:
            story.append(Paragraph(
                f"• {cert.name} - {cert.issuer} ({cert.date})",
                body_style
            ))

    doc.build(story)
    logger.info(f"PDF exported successfully: {output_path}")


# ──────────────────────────────────────────────────────────────
# DOCX Export
# ──────────────────────────────────────────────────────────────


def export_to_docx(resume: Resume, output_path: Path) -> None:
    """
    Export resume to DOCX using python-docx.

    Args:
        resume: Resume data.
        output_path: Destination file path.
    """
    logger.info(f"Exporting resume to DOCX: {output_path}")

    doc = Document()

    # Personal Info
    pi = resume.personal_info
    name_para = doc.add_paragraph(pi.full_name)
    name_para.runs[0].font.size = Pt(20)
    name_para.runs[0].font.bold = True
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_para = doc.add_paragraph(
        f"{pi.email} | {pi.phone} | {pi.location}"
    )
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if pi.linkedin or pi.github:
        links = []
        if pi.linkedin:
            links.append(f"LinkedIn: {pi.linkedin}")
        if pi.github:
            links.append(f"GitHub: {pi.github}")
        link_para = doc.add_paragraph(" | ".join(links))
        link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Summary
    if pi.summary:
        _add_section_heading(doc, "PROFESSIONAL SUMMARY")
        doc.add_paragraph(pi.summary)

    # Experience
    if resume.experience:
        _add_section_heading(doc, "EXPERIENCE")
        for exp in resume.experience:
            job_para = doc.add_paragraph()
            job_para.add_run(f"{exp.position}").bold = True
            job_para.add_run(f" at {exp.company}")

            doc.add_paragraph(
                f"{exp.location} | {exp.start_date} - {exp.end_date}"
            )

            for item in exp.responsibilities + exp.achievements:
                doc.add_paragraph(item, style="List Bullet")

            doc.add_paragraph()  # Spacing

    # Education
    if resume.education:
        _add_section_heading(doc, "EDUCATION")
        for edu in resume.education:
            degree_para = doc.add_paragraph()
            degree_para.add_run(f"{edu.degree} in {edu.field}").bold = True

            doc.add_paragraph(
                f"{edu.institution} | {edu.start_date} - {edu.end_date}"
            )
            if edu.gpa:
                doc.add_paragraph(f"GPA: {edu.gpa}")

    # Skills
    if resume.skills:
        _add_section_heading(doc, "SKILLS")
        doc.add_paragraph(", ".join(resume.skills))

    # Projects
    if resume.projects:
        _add_section_heading(doc, "PROJECTS")
        for proj in resume.projects:
            proj_para = doc.add_paragraph()
            proj_para.add_run(proj.title).bold = True
            doc.add_paragraph(proj.description)
            if proj.technologies:
                doc.add_paragraph(
                    f"Technologies: {', '.join(proj.technologies)}"
                )

    # Certifications
    if resume.certifications:
        _add_section_heading(doc, "CERTIFICATIONS")
        for cert in resume.certifications:
            doc.add_paragraph(
                f"• {cert.name} - {cert.issuer} ({cert.date})"
            )

    doc.save(str(output_path))
    logger.info(f"DOCX exported successfully: {output_path}")


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a formatted section heading to a Word document."""
    para = doc.add_paragraph(text)
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True
    para.runs[0].font.color.rgb = RGBColor(52, 73, 94)  # Dark blue-gray


# ──────────────────────────────────────────────────────────────
# TXT Export
# ──────────────────────────────────────────────────────────────


def export_to_txt(resume: Resume, output_path: Path) -> None:
    """
    Export resume to plain text format.

    Args:
        resume: Resume data.
        output_path: Destination file path.
    """
    logger.info(f"Exporting resume to TXT: {output_path}")

    lines = []
    pi = resume.personal_info

    # Header
    lines.append("=" * 60)
    lines.append(pi.full_name.center(60))
    lines.append(f"{pi.email} | {pi.phone} | {pi.location}".center(60))
    if pi.linkedin:
        lines.append(f"LinkedIn: {pi.linkedin}".center(60))
    lines.append("=" * 60)
    lines.append("")

    # Summary
    if pi.summary:
        lines.append("PROFESSIONAL SUMMARY")
        lines.append("-" * 60)
        lines.append(pi.summary)
        lines.append("")

    # Experience
    if resume.experience:
        lines.append("EXPERIENCE")
        lines.append("-" * 60)
        for exp in resume.experience:
            lines.append(f"{exp.position} at {exp.company}")
            lines.append(f"{exp.location} | {exp.start_date} - {exp.end_date}")
            for item in exp.responsibilities + exp.achievements:
                lines.append(f"  • {item}")
            lines.append("")

    # Education
    if resume.education:
        lines.append("EDUCATION")
        lines.append("-" * 60)
        for edu in resume.education:
            lines.append(f"{edu.degree} in {edu.field}")
            lines.append(f"{edu.institution} | {edu.start_date} - {edu.end_date}")
            if edu.gpa:
                lines.append(f"GPA: {edu.gpa}")
            lines.append("")

    # Skills
    if resume.skills:
        lines.append("SKILLS")
        lines.append("-" * 60)
        lines.append(", ".join(resume.skills))
        lines.append("")

    # Projects
    if resume.projects:
        lines.append("PROJECTS")
        lines.append("-" * 60)
        for proj in resume.projects:
            lines.append(proj.title)
            lines.append(proj.description)
            if proj.technologies:
                lines.append(f"Technologies: {', '.join(proj.technologies)}")
            lines.append("")

    # Certifications
    if resume.certifications:
        lines.append("CERTIFICATIONS")
        lines.append("-" * 60)
        for cert in resume.certifications:
            lines.append(f"• {cert.name} - {cert.issuer} ({cert.date})")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info(f"TXT exported successfully: {output_path}")
